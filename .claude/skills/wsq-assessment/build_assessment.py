#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Design Thinking Course for Businesses' (TGS-2026064719).

This is a REVISION of the assessment on file in the course's Google Drive, not a redesign.
It mirrors the original exactly:

  Original (Drive → Assessment/):
    · "Oral Questioning (OQ) - Design Thinking Course for Businesses - v1"   → 7 questions, K1–K7
    · "PP Assessment - Design Thinking Course for Businesses - v6"           → 4 tasks, A1–A6, 70 min

  This revision keeps: the same TWO instruments, the SAME question/task counts (7 and 4),
  the SAME K/A codes and mapping, and the SAME timings. Only the CONTENT is rewritten —
  fresh scenario, questions and model answers drawn from this course's slides and labs.

Produces four DOCX (question paper + answer key for each instrument), all with the WSQ
house cover page. Page 1 is the cover; page 2 carries Trainee Information + Instructions +
Grading; the questions/tasks begin on page 3. Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
for _cand in (os.path.join(REPO, ".claude/skills/courseware-build/build"),
              os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Design Thinking Course for Businesses"
COURSE_CODE = "TGS-2026064719"
TSC         = "Design Thinking Practices (DSN-ACE-3014-1.1)"
OQ_MINUTES  = "20 minutes"          # per the approved assessment plan
PP_MINUTES  = "70 minutes"          # unchanged from the original PP v6
# ────────────────────────────────────────────────────────────────────────────
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT = os.path.join(REPO, "assessment")
os.makedirs(OUT, exist_ok=True)

def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, ".claude/skills/courseware-build/assets", name),
              os.path.join(REPO, "courseware/assets", name),
              os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = None

Q_VER, A_VER = "v7", "v7"
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE) — 7 questions, K1–K7
# Mirrors the original OQ v1 question-for-question and code-for-code.
# (codes, context, question, [model-answer points])
WRITTEN = [
 ("K1, K2, K3, K6",
  "Design thinking is described throughout this course as a human-centred approach to problem solving that "
  "aims to improve people's experience — as much a mindset as a process.",
  "What are the key concepts of design thinking?",
  ["It is a problem-solving approach that aims to improve people's experience, not merely to ship a product.",
   "It is human-centred: it starts from a deep understanding of customers' needs and wants.",
   "It encourages creative consideration of a wide array of innovative solutions rather than the first obvious one.",
   "It is as much a mindset as a process — how you frame the problem matters as much as the steps you follow.",
   "It is iterative: the five action phases form a loop, not a straight line.",
   "It deliberately combines analytical (left-brain) and creative (right-brain) thinking.",
   "For an organisation it sits at the intersection of desirability (customers), viability (business) and "
   "feasibility (technology). (Slides: 'What is Design Thinking?', 'Design Thinking for the Organisation')"]),
 ("K1, K2, K3, K6",
  "The Design Management Institute tracked design-led organisations against the wider market to quantify what "
  "design thinking is actually worth to a business.",
  "What are the benefits of design thinking to a business?",
  ["It offers a structured problem-solving roadmap instead of ad-hoc brainstorming.",
   "It helps unlock innovative ideas that a business-centric or technology-first approach would miss.",
   "It solves genuine human problems, so solutions are adopted rather than resisted.",
   "It creates behavioural change and increases customer satisfaction and loyalty.",
   "It gains competitive advantage and reduces the risk of building the wrong thing.",
   "Measured returns (Design Management Institute): 41% higher market share, 46% greater competitive advantage "
   "overall, 50% more loyal customers, and 70% reporting that their digital experiences beat competitors. "
   "(Slides: 'Benefits of Design Thinking', 'The ROI of Design Thinking')"]),
 ("K1, K2, K3, K6",
  "Design thinking is a mindset anyone can apply, not a skill reserved for professional designers. The course "
  "contrasts the traditional thinker with the design thinker.",
  "What are the traits and the mindset of a design thinker?",
  ["Think users first — start from the human being served, not from the business or the technology.",
   "Ask the right questions — reframe a narrow brief into a human-centred one before solving it.",
   "Believe you can sketch — make ideas visible; you do not need to be an artist.",
   "Commit to explore — go wide and defer judgement before narrowing.",
   "Prototype to test and evaluate — build to think, and let evidence decide.",
   "Assume a beginner's mindset — consciously set aside your own assumptions and expertise when observing users.",
   "Hold a growth mindset — ability is a starting point, not a ceiling; keep iterating rather than defending "
   "the first idea.",
   "Contrast with the traditional thinker: 'we have this problem, let's brainstorm solutions', 'our competitors "
   "launched X, how do we copy it quickly?', 'we have this technology, what can we use it for?'. "
   "(Slides: 'Mindset — Traditional vs Design Thinker', 'Design thinking is a growth mindset')"]),
 ("K1, K2, K3, K6",
  "The problems an enterprise faces are much larger than designing a single product, so design thinking has to "
  "be embedded deliberately rather than run once as a workshop.",
  "What are the methods of applying and embedding design thinking in your organization?",
  ["Start by finding friction — where customers or staff visibly struggle, complain or invent workarounds.",
   "Treat repeated complaints, support tickets and journey drop-off points as free user research.",
   "Ask the frontline: the people serving customers already know where the pain is.",
   "Rank the opportunities by impact and effort, and start with one well-chosen initiative rather than many.",
   "Run a short design sprint on that opportunity: Empathize and Define, then Ideate, Prototype and Test.",
   "Test every idea against the three lenses — desirability, viability and feasibility.",
   "Give the concept a clear owner (for example using a RACI matrix) so it survives the handover to delivery.",
   "Define the success metrics before building, and build a culture of experimentation that tolerates "
   "informative failure. (Slides: 'Uncovering Opportunities in Your Organisation', 'Take It Back to Work')"]),
 ("K5",
  "Design thinking has been applied across corporate, service and public-sector organisations.",
  "Name one use case of design thinking in another organization and explain what changed as a result.",
  ["Any ONE of the cases covered in class, with the user, the insight and the outcome:",
   "Airbnb — built a culture of experimentation and design-led iteration, going out to meet hosts and "
   "re-photographing listings by hand; moved from a failing start-up to a billion-dollar business.",
   "IBM — applies IBM Design Thinking at scale, restructuring how large, complex product teams work.",
   "Bank of America — partnered with IDEO in 2004; observing real people revealed that savers round up their "
   "spending, producing the 'Keep the Change' account.",
   "Uber Eats — designs its delivery service around the observed journeys of customers, couriers and restaurants.",
   "Healthcare (Stanford Hasso Plattner Institute) — redesigned the emergency-room patient experience.",
   "Clean Team — in-home toilets for Ghana's urban poor; Golden Gate Regional Center — redesigned disability "
   "services; The Good Kitchen (Denmark) — a municipal meal service redesigned for quality, flexibility and "
   "choice. (Slides: 'Design Thinking in Business', 'Design Thinking in Services & Society')"]),
 ("K4",
  "The five-stage model was proposed by the Hasso-Plattner Institute of Design at Stanford University (d.school).",
  "What are the stages in the design thinking process, and what is the objective of each?",
  ["EMPATHIZE — understand the experience, situation and emotion of the user you are designing for, by "
   "observing, engaging and immersing, without judging.",
   "DEFINE — process and synthesise the findings into a meaningful, actionable problem statement "
   "(user + need + insight).",
   "IDEATE — go wide: generate a large quantity and variety of ideas to get beyond the obvious solution.",
   "PROTOTYPE — build to think: a simple, cheap, fast artefact that makes the idea tangible.",
   "TEST — put the prototype in users' hands, gather feedback, then refine the prototype or reframe the problem.",
   "The five stages form a LOOP, not a line — insights from Test routinely send you back to Empathize or "
   "Define. Iteration is the method, not a sign of failure. (Slides: 'The Five Action Phases')"]),
 ("K7",
  "A prototype exists to provoke a reaction, not to impress. It should be quick and cheap enough that you are "
  "willing to throw it away.",
  "Give some examples of how to create prototypes, and explain the difference between low- and high-fidelity "
  "prototyping.",
  ["A prototype can take any physical form: a wall of post-it notes, a role-play, a space, an object, an "
   "interface or a storyboard.",
   "Low fidelity — paper sketches, storyboards, cardboard, foam, clay or building blocks. Fast, cheap and easy "
   "to change; used early to test the concept; invites honest criticism because it clearly is not finished.",
   "High fidelity — wireframes, interactive mock-ups and near-final materials. Slower and more expensive; used "
   "later to test detail; carries the risk that people critique the polish rather than the idea.",
   "Guidelines: just start building, don't spend too much time, remember what you are testing for, and build "
   "with the user in mind.",
   "Tools used in class include paper and sticky notes, draw.chat and the Prototype canvas in the Design "
   "Thinking Toolkit (https://alfredang.github.io/designthinking/). "
   "(Slides: 'Action Phase 4 — Prototype', 'Low- vs High-Fidelity Prototyping')"]),
]

# ---------------------------------------------------------------- PRACTICAL (ABILITY) — 4 tasks, A1–A6
SCENARIO = (
    "MediCare Family Clinic is a busy neighbourhood clinic in Singapore. Patients routinely wait 45–90 minutes "
    "past their appointment time, and the waiting room is regularly full. Patients complain that nobody tells "
    "them how long the wait will be; the front-desk staff are interrupted constantly by patients asking when "
    "they will be seen; and the doctors are frustrated because consultations are rushed to catch up on the "
    "backlog. Online reviews have fallen sharply, and the clinic manager has engaged you as a consultant.\n\n"
    "The manager's instinct is to 'just hire another receptionist'. You have been asked instead to apply a "
    "design thinking approach so that the clinic solves the right problem. Complete the four tasks below, "
    "working through the action phases in order."
)

PRACTICAL = [
 ("Task 1", "A1",
  "First you need to know who your users are and understand their current experience. Your users and "
  "stakeholders are the patients, the front-desk staff and the doctors. You have observed the waiting room "
  "and interviewed people from each group.\n\n"
  "Create an EMPATHY MAP below for ONE of these stakeholder groups. Capture what they SAY (use their own "
  "words), what they THINK but do not say out loud, what they DO that you can observe, and what they FEEL and "
  "how strongly. Then state ONE insight you found — in particular any contradiction between the quadrants, "
  "which is usually where the real insight hides.",
  "Complete your empathy map in the box below. (Taught in Activity 5 — Empathy Map & Persona (A1, A5, K7), "
  "and applied in Activity 4 — The Wallet Project (A1, A4, K4). Ed-tool: "
  "https://alfredang.github.io/designthinking/)",
  "Suggestive answer for PATIENTS (not exhaustive — any one stakeholder group, mapped properly, is acceptable):\n\n"
  "SAYS (verbatim quotes)\n"
  "  \"My appointment was at 10, it's already 11:15.\"\n"
  "  \"Nobody tells you anything — you just sit there.\"\n"
  "  \"I had to take half a day off work for a 10-minute consult.\"\n\n"
  "THINKS (not said out loud)\n"
  "  \"Have they forgotten me? Did I miss my name being called?\"\n"
  "  \"If I go to the toilet or step out for a drink I'll lose my turn.\"\n"
  "  \"Am I going to be late picking up my child?\"\n\n"
  "DOES (observable)\n"
  "  Repeatedly looks up at the queue display and at the consultation-room door.\n"
  "  Walks up to the front desk to ask how many people are ahead — often more than once.\n"
  "  Does not leave the room, even to get food or drink.\n\n"
  "FEELS (emotion + intensity)\n"
  "  Anxious and powerless — high. Frustrated — high. Resigned — moderate.\n\n"
  "INSIGHT / CONTRADICTION\n"
  "The complaint is about the LENGTH of the wait, but behaviour shows the real problem is the UNCERTAINTY of "
  "it: patients stay rooted to their seats because they cannot risk missing their turn. A patient who KNEW "
  "they had 40 minutes could get a coffee and would report a far better experience even with the same wait. "
  "The clinic does not have a speed problem so much as an information problem.\n\n"
  "Assessor note: award Competent where all four quadrants are populated with that group's own language, the "
  "quadrants are used correctly (SAYS is not THINKS, DOES is observable, FEELS is an emotion), and at least "
  "one genuine insight or contradiction is identified rather than a restatement of the complaint."),
 ("Task 2", "A2",
  "Having gained insight into your stakeholder's experience through empathising, now DEFINE the problem.\n\n"
  "Write a good problem statement using the POINT OF VIEW (POV) template:\n\n"
  "    [USER] needs a way to [USER'S NEED] because [INSIGHT].\n\n"
  "Use the same stakeholder group you mapped in Task 1. Remember that the need must be a VERB and the insight "
  "must be a genuine discovery, not a restatement of the need. Then list any underlying assumptions your "
  "statement makes, and reframe the POV as ONE 'How Might We' question.",
  "Write your POV statement, assumptions and HMW question in the box below. (Taught in Activity 6 — POV, How "
  "Might We & Brainstorming (A1, A2, A3, A4, K7); opportunity framing from Activity 3 (A2, K5). Ed-tool: "
  "https://alfredang.github.io/designthinking/)",
  "Suggestive answer (not exhaustive):\n\n"
  "POV STATEMENT\n"
  "  A working patient at MediCare Family Clinic needs a way to know how long their wait will actually be and "
  "to reclaim that time, because the uncertainty — not the delay itself — is what traps them in the waiting "
  "room and makes the visit feel like a lost half-day.\n\n"
  "  USER    — a working patient with a scheduled appointment\n"
  "  NEED    — to know (verb) how long the wait will be and to use that time\n"
  "  INSIGHT — the uncertainty, not the delay, is what causes the frustration\n\n"
  "UNDERLYING ASSUMPTIONS\n"
  "  · That the clinic can estimate the remaining wait with reasonable accuracy.\n"
  "  · That most patients carry a mobile phone and would act on a notification.\n"
  "  · That patients would genuinely leave the waiting room if they trusted the estimate.\n"
  "  · That an approximate estimate is better received than no information at all.\n\n"
  "HOW MIGHT WE\n"
  "  How might we give patients enough confidence about their waiting time that they can spend it as they choose?\n\n"
  "Assessor note: award Competent where the statement follows the USER + NEED + INSIGHT structure, the need is "
  "expressed as a verb, the insight is a discovery drawn from Task 1 rather than a repeat of the need, "
  "assumptions are stated, and the HMW is open enough to admit many answers while still giving direction."),
 ("Task 3", "A3, A4",
  "With a good problem statement you can gather a diverse team and brainstorm. Over the course of a brainstorm "
  "ideas ebb and flow — they start boring and move between absurd and brilliant. The more absurd the idea, the "
  "more likely it is to spawn a brilliant one.\n\n"
  "Run a brainstorm against your 'How Might We' question from Task 2. Record AT LEAST SIX ideas, including at "
  "least one deliberately wild one. Then CONVERGE: select the three strongest ideas using a stated selection "
  "method (for example dot voting or a Now-Wow-How matrix), and explain WHY each of the three was chosen.",
  "List your ideas, your selection method and your shortlist in the box below. (Taught in Activity 6 — POV, "
  "How Might We & Brainstorming (A1, A2, A3, A4, K7); idea selection also practised in Activity 7 (A3, A5, A6, "
  "K7). Ed-tool: https://alfredang.github.io/designthinking/)",
  "Suggestive answer (not exhaustive):\n\n"
  "IDEAS GENERATED (diverge — quantity over quality, judgement deferred)\n"
  "  1. SMS or WhatsApp alert sent when the patient is three places from being called.\n"
  "  2. A live queue-position display, visible on a screen and on the patient's phone.\n"
  "  3. A 'go and come back' pass — leave the clinic, keep your place, return when notified.\n"
  "  4. Realistic appointment slots based on the actual average consultation time, not the ideal one.\n"
  "  5. Triage at check-in so short visits (repeat prescriptions, MC collection) are routed to a fast lane.\n"
  "  6. Partner with the coffee shop next door — show your queue number for a discount while you wait.\n"
  "  7. (Wild) Deliver the consultation by video for anyone whose case does not need a physical examination.\n"
  "  8. (Wild) Pay patients for every 15 minutes the clinic runs late.\n\n"
  "SELECTION METHOD\n"
  "  Now-Wow-How matrix, scoring each idea on impact for the patient against effort for the clinic; the team "
  "then dot-voted within the top band.\n\n"
  "SHORTLIST AND RATIONALE\n"
  "  · Idea 3 — 'Go and come back' pass (NOW). Directly answers the HMW: it converts dead waiting time into the "
  "patient's own time. Low cost, and it addresses the insight rather than the symptom.\n"
  "  · Idea 1 — SMS/WhatsApp alert (NOW). The enabling mechanism that makes idea 3 trustworthy; uses "
  "infrastructure the clinic already has.\n"
  "  · Idea 5 — Fast lane triage (WOW). Attacks the root cause of the backlog rather than the experience of it; "
  "higher effort, but the largest long-term reduction in average wait.\n\n"
  "Assessor note: award Competent where at least six genuinely different ideas are recorded (including a wild "
  "one), a named selection method is applied, exactly three ideas are shortlisted, and each carries a stated "
  "reason linked back to the POV or HMW rather than personal preference."),
 ("Task 4", "A5, A6",
  "Select ONE of the shortlisted ideas from Task 3 and take it through Prototype and Test.\n\n"
  "Part A — Describe or draw a LOW-FIDELITY PROTOTYPE of your chosen idea. State what it is made of and what "
  "it lets the user do.\n\n"
  "Part B — State the single most important ASSUMPTION your idea depends on, and describe the cheapest, "
  "fastest test that could disprove it.\n\n"
  "Part C — You have shown the prototype to patients and staff. Record the feedback you would expect — what "
  "worked and what confused them — and state how you would IMPROVE the prototype in response.\n\n"
  "Part D — State the METRICS you would use to prove the solution worked, and explain how you would COMMUNICATE "
  "the outcome and its value to the clinic manager.",
  "Complete Parts A–D in the box below. (Taught in Activity 7 — Prototype, Test & Measure the Outcome "
  "(A3, A5, A6, K7): Parts A–C follow the prototype and test steps, and Part D follows the metrics and "
  "stakeholder-communication steps of that activity. Ed-tools: https://alfredang.github.io/designthinking/ "
  "and https://alfredang.github.io/raci/)",
  "Suggestive answer (not exhaustive):\n\n"
  "PART A — LOW-FIDELITY PROTOTYPE\n"
  "A paper 'Go & Come Back' pass, card-sized, filled in by hand at check-in: queue number, estimated "
  "call-back time, and the words 'We will text you when you are 3 patients away.' Paired with a paper "
  "storyboard of the four SMS messages the patient would receive. Total build cost: a stack of cards and a "
  "marker pen.\n\n"
  "PART B — TOP ASSUMPTION AND THE CHEAPEST TEST\n"
  "Assumption: patients will actually leave the waiting room if they trust the estimate.\n"
  "Cheapest test: for one afternoon, hand the paper pass to 15 patients and have a staff member send the SMS "
  "manually from a mobile phone. Measure how many leave the waiting room and how many return late or miss "
  "their turn. Cost: one afternoon, no software. Success threshold: at least half leave and none miss a turn.\n\n"
  "PART C — EXPECTED FEEDBACK AND REFINEMENT\n"
  "Worked — patients liked being given permission to leave; several said it was the first time the clinic had "
  "told them anything; front-desk interruptions dropped noticeably.\n"
  "Confused — 'How far can I actually go?' and 'What if I don't hear the phone?'; elderly patients without a "
  "mobile phone were excluded; some did not trust the estimate the first time.\n"
  "Improve — print a suggested radius ('stay within 10 minutes' walk') on the pass; add a second alert at 5 "
  "patients away; keep a physical roll-call fallback for patients without a phone; show a RANGE (e.g. 35–50 "
  "min) rather than a single time so the estimate stays credible.\n\n"
  "PART D — METRICS AND COMMUNICATION\n"
  "Traditional KPI — average patient wait time; number of consultations per session; no-show rate.\n"
  "Customer feedback — patient satisfaction score, Net Promoter Score, average online review rating.\n"
  "Activity metrics — number of passes issued; percentage of patients who leave and return on time.\n"
  "Quick results — front-desk interruptions per hour (before vs after); staff-reported stress.\n"
  "Value / novelty — plot the concept on the business-value versus novelty grid; the pass is valuable and "
  "novel for this clinic, so it is worth resourcing.\n"
  "Communicating it — present to the manager as before/after evidence: the original complaint, the insight "
  "that uncertainty rather than delay drove it, the prototype, the test result and the metric movement. Note "
  "that the original instinct — hiring another receptionist — would have added cost without addressing the "
  "insight. Assign ownership with a RACI matrix (https://alfredang.github.io/raci/) so the rollout has an "
  "accountable name against it.\n\n"
  "Assessor note: award Competent where all four parts are completed; the prototype is genuinely low-fidelity "
  "and tangible; a single testable assumption is named with a cheap test and a success threshold; the "
  "refinement is a direct response to the stated feedback; and the metrics span more than one category with a "
  "clear explanation of how the outcome and its value would be communicated to the stakeholder."),
]

# ---------------------------------------------------------------- coverage check
def check_coverage():
    """Fail the build if any K is unassessed by the WA or any A unassessed by the PP."""
    import re
    def codes(rows, idx, letter):
        found = set()
        for r in rows:
            for c in re.findall(rf"{letter}\d+", r[idx]):
                found.add(c)
        return found
    ks = codes(WRITTEN, 0, "K")
    as_ = codes(PRACTICAL, 1, "A")
    want_k = {f"K{i}" for i in range(1, 8)}
    want_a = {f"A{i}" for i in range(1, 7)}
    print("  WA covers:", ", ".join(sorted(ks, key=lambda x: int(x[1:]))))
    print("  PP covers:", ", ".join(sorted(as_, key=lambda x: int(x[1:]))))
    missing_k, missing_a = want_k - ks, want_a - as_
    if missing_k or missing_a:
        raise SystemExit(f"COVERAGE FAILURE — missing K: {sorted(missing_k)}  missing A: {sorted(missing_a)}")
    print("  coverage OK — all K1–K7 and A1–A6 assessed.")

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet model answer; `code` → preformatted block;
    neither → empty answer space for the candidate."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    # Keep the whole box on one page — a bordered answer box cut by a page break is a
    # formatting failure in the house standard.
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))
    if code:
        # Model answers for this course are PROSE, not code. Rendering them in a fixed
        # monospace grid guarantees ugly wrapping (a wrapped line returns to column 0 and
        # destroys any alignment), so they are laid out as ordinary proportional text:
        # ALL-CAPS section labels are bolded, everything else flows and wraps naturally.
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for raw in code.split("\n"):
            ln = raw.rstrip()
            if not ln.strip():
                sp = cell.add_paragraph(style=None)
                sp.paragraph_format.space_after = Pt(0); sp.paragraph_format.space_before = Pt(0)
                sp.add_run(" ").font.size = Pt(4)
                continue
            stripped = ln.strip()
            indent = len(ln) - len(ln.lstrip())
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(1); b.paragraph_format.space_before = Pt(0)
            if indent:
                b.paragraph_format.left_indent = Inches(0.12 * min(indent // 2, 3))
            rr = b.add_run(stripped)
            rr.font.size = Pt(9.5)
            # A heading line is a short ALL-CAPS label such as "PART A — LOW-FIDELITY PROTOTYPE"
            letters = [c for c in stripped if c.isalpha()]
            if letters and all(c.isupper() for c in letters[:18]) and len(stripped) < 70:
                rr.bold = True
            elif stripped.startswith("Assessor note"):
                rr.italic = True
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text, oral=False):
    """Instructions to the candidate. For Oral Questioning the answers are spoken and
    recorded by the assessor, so the LMS-upload step does not apply."""
    heading(doc, "Instructions to Candidate")
    if oral:
        items = [
            "This is an individual assessment.",
            "This is an open-book assessment. You may refer to the course slides, the Learner Guide and "
            "approved materials only.",
            f"A total of {minutes_text} is given to complete this assessment.",
            "The assessor will ask you each question verbally and record your response. Answer in your own "
            "words and give an example wherever you can.",
            "The assessor may ask a follow-up question to confirm your understanding.",
        ] + BRIEFING[:3]
        for i, s in enumerate(items, 1):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            p.add_run(f"{i}.  {s}").font.size = Pt(11)
        return
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment. You may refer to the course slides, the Learner Guide and "
        "approved materials only.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    """The assessor sign-off block. It lives on PAGE 2 — never at the back of the paper."""
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    """Oral Questioning (OQ) — the assessor asks these verbally and records the response.
    The candidate does not write on this paper; the boxes are the ASSESSOR's record."""
    doc = base_doc()
    kind = "Oral Questioning (OQ) — Answer Key" if answers else "Oral Questioning (OQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO, course_code=COURSE_CODE)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Oral Questioning (OQ)" if answers else "Oral Questioning (OQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}  ·  TSC: {TSC}", size=10.5, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        candidate_block(doc); instructions(doc, OQ_MINUTES, oral=True)
        grading(doc, "Candidate has answered all oral questions and demonstrated the underpinning "
                     "knowledge (K1–K7) required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Oral Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "The assessor asks each question verbally and records the candidate's response in the box "
              "below it. Questions are open-ended — there are no multiple-choice options. The assessor may "
              "ask a follow-up question to confirm the candidate's understanding.",
         size=10.5, italic=True, color=GREY, after=8)
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=110)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answers to Oral Questioning (OQ) - {TITLE} - {suffix}.docx" if answers
            else f"Oral Questioning (OQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO, course_code=COURSE_CODE)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}  ·  TSC: {TSC}", size=10.5, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        candidate_block(doc); instructions(doc, PP_MINUTES)
        grading(doc, "Candidate has successfully completed all four practical tasks, demonstrating the "
                     "abilities (A1–A6) required for the course learning outcomes, and can explain the "
                     "design thinking methods and tools used to achieve them.")
        page_break(doc)
    para(doc, "Practical Performance", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # On the QUESTION PAPER task 1 shares the scenario page and each later task gets its
    # own page. In the ANSWER KEY every task starts on a fresh page — the model-answer
    # blocks are tall, and a box that straddles a page break is a formatting failure.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        if answers or i > 1:
            page_break(doc)
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=230)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    check_coverage()
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. OQ: {len(WRITTEN)} questions (K1–K7) · PP: {len(PRACTICAL)} tasks (A1–A6).")
