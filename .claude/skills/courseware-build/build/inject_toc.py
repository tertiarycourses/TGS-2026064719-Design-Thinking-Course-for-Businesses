#!/usr/bin/env python3
"""Replace the placeholder TOC in a DOCX with a static, page-numbered TOC.

Page numbers are read from the already-rendered PDF (same basename) so the
TOC in the re-rendered PDF shows real page numbers. Assumes the TOC occupies
a single page in both passes (true when entry count is small), so body page
numbers do not shift between passes.

Usage: python3 inject_toc.py <docx> <pdf> [maxlevel]
"""
import sys, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from pypdf import PdfReader

def norm(s):
    return re.sub(r'[^a-z0-9]', '', s.lower())

def main():
    docx_path, pdf_path = sys.argv[1], sys.argv[2]
    maxlevel = int(sys.argv[3]) if len(sys.argv) > 3 else 2

    doc = Document(docx_path)

    # 1) collect heading paragraphs (in order) up to maxlevel
    heads = []
    for p in doc.paragraphs:
        sn = p.style.name
        if sn.startswith("Heading"):
            try:
                lvl = int(sn.split()[-1])
            except ValueError:
                continue
            if lvl <= maxlevel and p.text.strip():
                heads.append((lvl, p.text.strip()))

    # 2) page text from the rendered PDF
    reader = PdfReader(pdf_path)
    pages = [norm(pg.extract_text() or "") for pg in reader.pages]

    # 3) map each heading -> first page (>= cursor) whose text contains it
    entries = []
    cursor = 0
    for lvl, text in heads:
        key = norm(text)[:24]
        page = None
        for i in range(cursor, len(pages)):
            if key and key in pages[i]:
                page = i + 1  # 1-based, matches footer "Page N"
                cursor = i
                break
        if page is None:  # fallback: search from start
            for i in range(len(pages)):
                if key and key in pages[i]:
                    page = i + 1
                    break
        entries.append((lvl, text, page or 1))

    # 3b) a multi-page static TOC pushes every body page down between passes
    #     (the pass-1 PDF was measured with the ~1-page placeholder). Offset by
    #     the extra pages the injected TOC occupies. offset is 0 for small TOCs
    #     (e.g. the Lesson Plan), so this never regresses single-page documents.
    LINES_PER_PAGE = 44
    toc_lines = len(entries) + 2            # "Contents" heading + spacing
    toc_pages = max(1, -(-toc_lines // LINES_PER_PAGE))
    offset = toc_pages - 1
    if offset:
        entries = [(lvl, text, pg + offset) for lvl, text, pg in entries]

    # 4) find the placeholder TOC paragraph (contains a TOC field or the hint text)
    placeholder = None
    for p in doc.paragraphs:
        xml = p._p.xml
        if ("Update Field" in p.text) or ("TOC " in xml and 'instrText' in xml) or ("fldSimple" in xml and "TOC" in xml):
            placeholder = p
            break
    if placeholder is None:
        print("  [inject_toc] no TOC placeholder found in", docx_path); return

    # 5) Build the TOC as a REAL Word TOC field whose cached result is the
    #    page-numbered static text. Word renders the cached result immediately
    #    (so the headless PDF is correct) and F9 / "Update Field" still works.
    #
    #    Structure:  [begin + instrText TOC]  entry paras…  [end]
    anchor = placeholder._p
    GREY = RGBColor(0x33, 0x33, 0x33)
    from docx.text.paragraph import Paragraph

    def _mkpara():
        new_p = anchor.makeelement(qn('w:p'), {})
        anchor.addprevious(new_p)
        return Paragraph(new_p, placeholder._parent)

    def _fld(para, kind, instr=None):
        r = para.add_run()
        fc = r._r.makeelement(qn('w:fldChar'), {})
        fc.set(qn('w:fldCharType'), kind)
        r._r.append(fc)
        if instr is not None:
            r2 = para.add_run()
            it = r2._r.makeelement(qn('w:instrText'), {})
            it.set(qn('xml:space'), 'preserve')
            it.text = instr
            r2._r.append(it)

    for i, (lvl, text, page) in enumerate(entries):
        para = _mkpara()
        pf = para.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        if lvl >= 2:
            pf.left_indent = Inches(0.3)
        pf.space_after = Pt(3)
        if i == 0:
            # open the field on the first entry paragraph
            _fld(para, 'begin', f'TOC \\o "1-{maxlevel}" \\h \\z \\u')
            _fld(para, 'separate')
        r = para.add_run(text + "\t" + str(page))
        r.font.size = Pt(11 if lvl == 1 else 10.5)
        r.font.name = "Arial"
        r.bold = (lvl == 1)
        r.font.color.rgb = GREY
        if i == len(entries) - 1:
            _fld(para, 'end')     # close the field on the last entry
    anchor.getparent().remove(anchor)

    doc.save(docx_path)
    print(f"  [inject_toc] {docx_path}: wrote {len(entries)} TOC entries "
          f"(pages {entries[0][2]}..{entries[-1][2]})")

if __name__ == "__main__":
    main()
