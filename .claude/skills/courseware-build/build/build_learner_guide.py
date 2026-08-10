#!/usr/bin/env python3
"""Generate the Design Thinking Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per activity (Objective · Goal · What you'll produce ·
Step-by-step · You're done when). The step-by-step detail lives HERE and in the
labs/ files — by design the slide deck carries none.
"""
import os, sys, copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  f"It provides step-by-step instructions for all {len(ACT)} hands-on activities, organised by the four "
  f"course topics, and is aligned to the Skills Framework for Design TSC \"{C.TSC_TITLE}\" ({C.TSC_CODE}).")
p("Design thinking is a human-centred approach to problem solving that aims to improve people's "
  "experience. It is as much a mindset as a process. This guide is written so you can follow every "
  "activity in class and repeat it afterwards at your own workplace — the activities use free, "
  "browser-based ed-tools that stay available to you after the course.")
p("Use this guide alongside the course slides and the activity files in the labs/ folder. The slide "
  "deck is deliberately visual and does not repeat the step-by-step instructions — those live here.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"TSC Title: {C.TSC_TITLE}   ·   TSC Code: {C.TSC_CODE}")
h3("Abilities")
bullets(C.TSC_ABILITIES)
h3("Knowledge")
bullets(C.TSC_KNOWLEDGE)

h1("Before You Start — Your Ed-Tools")
p("Every activity in this course runs in your web browser. Nothing needs to be installed, and every "
  "tool below remains free to use after the course, so you can run the same activities with your own "
  "team.")
bullets([f"{nm} — {url} — {desc}" for nm,url,desc in C.EDTOOLS])
h3("Also used in class")
bullets([
 "draw.chat — a free collaborative whiteboard for quick sketching; share the board URL with the class.",
 "padlet.com — a shared wall used to post your work so the whole class can see and comment on it.",
 "Paper, pens and sticky notes — the fastest prototyping tools ever invented. Do not underestimate them.",
])
h3("Conventions used in every activity")
bullets([
 "Each activity states its objective, what you will produce, the steps, and how to check you are done.",
 "Timings shown are the in-class timings; take longer when you repeat the activity at work.",
 "Where an activity says 'interview your partner', use their exact words — verbatim quotes carry the insight.",
 "Do not skip the 'You're done when' check; it is the same evidence the trainer looks for in the practical assessment.",
])

h1("How to Download Your Course Material")
steps([
 ("Go to the LMS/TMS portal at https://lms-tms.tertiaryinfotech.com and sign in with the e-mail address you registered with.",""),
 ("Open My Courses and select Design Thinking Course for Businesses.",""),
 ("Download the Trainer Slides (PPT), Learner Slides (PDF), this Learner Guide (PDF) and the Lesson Plan (PDF).",""),
 ("Keep the slides and this guide open during the assessment — it is open book.",""),
 ("Complete the mandatory TRAQOM survey on the same portal at the end of the course.",""),
])

# ---------------- per-topic, per-activity ----------------
TOPIC_NOTES = {
 1: [
  ("What design thinking is",
   "Design thinking is a problem-solving approach that aims to improve people's experience. It is human-centred: it "
   "starts from a deep understanding of customers' needs and wants, encourages creative consideration of a wide array "
   "of innovative solutions, and is as much a mindset as a process."),
  ("Why it matters to a business",
   "Research by the Design Management Institute found design-led organisations reported 41% higher market share, 46% "
   "greater competitive advantage overall, 50% more loyal customers, and 70% saying their digital experiences beat "
   "competitors. Design thinking offers a problem-solving roadmap, unlocks innovative ideas, solves genuine human "
   "problems, creates behavioural change and increases customer satisfaction."),
  ("It is not just for designers",
   "Design thinking is a mindset anyone can apply. It means thinking like a designer so that you can create a new "
   "product or service that improves your customer's or user's experience — regardless of your job title."),
  ("Analytical and creative thinking together",
   "Left-brain thinking is analytical, rational, objective, focused on the present and past, on facts, order and "
   "planning. Right-brain thinking is creative, holistic, subjective, focused on the present and future, on feelings, "
   "space and spontaneity. Design thinking deliberately uses both."),
  ("The mindset shift",
   "Traditional thinkers say \"we have this problem, let's get in a room and brainstorm solutions\", \"our competitors "
   "just launched X, how can we do X quickly?\" or \"we have this technology, what can we use it for?\". Design thinking "
   "shifts focus from a business-centric engineering solution to a customer-centric one: think users first, ask the "
   "right questions, believe you can sketch, commit to explore, and prototype to test and evaluate. It is also a growth "
   "mindset — ability is a starting point, not a ceiling."),
  ("Applying it to an organisation",
   "The problems an enterprise faces are much bigger than designing a vase. Organisational design thinking happens at "
   "the intersection of three lenses: desirability for customers, viability at the business level, and feasibility for "
   "technology. A solution must satisfy all three."),
 ],
 2: [
  ("Airbnb",
   "Design thinking is part of Airbnb's success. In particular they built a culture of experimentation — going out to "
   "meet hosts, re-photographing listings by hand, and iterating relentlessly — which took them from a failing start-up "
   "to a billion-dollar business."),
  ("IBM",
   "IBM applies design thinking to complex teams, problems and organisations through its own IBM Design Thinking "
   "framework, restructuring how large product teams work."),
  ("Bank of America",
   "Bank of America partnered with design consultancy IDEO in 2004 to understand how to get more people to open bank "
   "accounts. Observing real people revealed that savers rounded up their spending — producing the 'Keep the Change' "
   "account."),
  ("Uber Eats",
   "The Uber Eats team designs its food-delivery service with a design-thinking mindset, studying the journeys of "
   "customers, couriers and restaurants in the field."),
  ("Healthcare — Hasso Plattner Institute",
   "The Hasso Plattner Institute of Design at Stanford explored design-thinking approaches to improve the patient "
   "experience in the emergency room."),
  ("Public and social sector",
   "Clean Team applied design thinking to provide in-home toilets for Ghana's urban poor. The Golden Gate Regional "
   "Center redesigned services and financial support for people with developmental disabilities. In Denmark, Hatch & "
   "Bloom redesigned a municipal meal service into The Good Kitchen, giving elderly residents more quality, flexibility "
   "and choice."),
  ("Uncovering your own opportunities",
   "Look for friction — where customers or staff visibly struggle, complain or invent workarounds. A workaround is a "
   "user telling you the design is wrong. Follow repeated complaints and support tickets; they are free user research. "
   "Find where people abandon a journey and ask what happened just before. Then rank the opportunities by impact and "
   "effort before committing a team."),
 ],
 3: [
  ("The five-stage model",
   "The five-stage design thinking model was proposed by the Hasso-Plattner Institute of Design at Stanford University "
   "(d.school): Empathize, Define, Ideate, Prototype and Test. The phases form a loop, not a straight line — insights "
   "from testing routinely send you back to empathising or redefining the problem."),
  ("Phase 1 — Empathize",
   "Design thinking cannot begin without a deep understanding of the people you are designing for. Objective: to "
   "understand the experience, situation and emotion of your user. Observe (view users and their behaviour in the "
   "context of their lives, without judging), Engage (interact in conversations and interviews; ask why) and Immerse "
   "(experience what your user experiences). Assume a beginner's mindset — consciously set aside your own assumptions. "
   "The classic example is the MRI scanner: immersion revealed that an MRI room is a frightening experience for a "
   "child, and the redesign turned the scan into an adventure story."),
  ("Phase 2 — Define",
   "An integral part of the process is defining a meaningful, actionable problem statement. Objective: process and "
   "synthesise your findings to form a problem statement you will address. Understand the User (the persona), the "
   "Needs (always expressed as verbs) and the Insights (what you discovered). A good problem statement is "
   "human-centred, broad enough to allow creativity, and narrow enough to be manageable. Methods include clustering "
   "and bundling ideas and facts, empathy mapping, POV statements, 'How Might We' questions and why-how laddering."),
  ("Phase 3 — Ideate",
   "Ideation is where you concentrate on idea generation — going wide in terms of concepts and outcomes. Objective: "
   "translate problems into solutions by exploring a wide variety and large quantity of ideas, going beyond the "
   "obvious. It relies on creativity (combining rational thought with imagination), group synergy (building on each "
   "other's ideas) and the separation of divergent from convergent thinking. Active facilitation matters: create a "
   "curious, courageous and concentrated atmosphere."),
  ("Phase 4 — Prototype",
   "Prototyping produces an early, inexpensive, scaled-down version of the product to reveal problems with the current "
   "design. Objective: build to think — a simple, cheap, fast way to shape ideas so you can experience and interact "
   "with them. A prototype can take any physical form: a wall of post-it notes, a role-play, a space, an object, an "
   "interface or a storyboard. Low-fidelity prototypes use basic models and simple materials; high-fidelity prototypes "
   "are closer to the finished product. Guidelines: just start building, don't spend too much time, remember what you "
   "are testing for, and build with the user in mind."),
  ("Phase 5 — Test",
   "Testing generates user feedback on your prototypes and deepens your understanding of your users. Objective: ask for "
   "feedback, learn about your user, reframe your POV and refine the prototype. Show (let people use it and listen), "
   "Create experiences (let them describe how it feels) and Compare (let users test multiple prototypes to reveal "
   "latent needs). Test in a natural setting wherever possible. Negative feedback is valuable: if users struggle, "
   "revisit your solutions and look for problems you had not considered."),
 ],
 4: [
  ("Methods by action phase",
   "Empathize produces empathy maps, persona maps, lists of user feedback and identified problems. Define produces a "
   "design brief (POV and HMW), stakeholder maps, context maps, customer journeys and opportunity maps. Ideate produces "
   "ideas and concepts, sketches, prioritisation maps, affinity maps and idea evaluations. Prototype produces physical "
   "prototypes, wireframes and storyboards. Test produces user feedback, observations, evaluation maps and proposed "
   "refinements."),
  ("The empathy map",
   "Empathy maps are split into four quadrants — Says, Thinks, Does and Feels — with the user or persona in the middle. "
   "SAYS contains verbatim quotes from the interview. THINKS captures what the user was thinking but did not say out "
   "loud. DOES records observable actions. FEELS records the emotional state and its intensity. Empathy maps give a "
   "glance into who a user is as a whole; they are not chronological. Users are complex, so contradictions between "
   "quadrants are normal and extremely beneficial — a positive quote next to a negative emotion is exactly where the "
   "real insight hides."),
  ("Persona and journey mapping",
   "Persona mapping identifies who your clients are and how they make decisions, then applies that to more effective "
   "strategies. Journey mapping charts every touchpoint over time — before, during and after the core interaction — "
   "exposing where the experience breaks down and which moments disproportionately shape the whole experience."),
  ("Point of View (POV)",
   "POV = USER + NEED + INSIGHT. The template is: [USER] needs a way to [USER'S NEED] because [INSIGHT]. For example: "
   "\"A busy working mom needs to be able to bring plastic bags with her when she's not at home, because she buys the "
   "equivalent of three dog-poop bags a day cleaning up after her dog.\" Or: \"A teacher needs to clean and store his "
   "used sandwich bags at school, because every morning he uses a new bag and often discards them rather than washing "
   "them.\" Developing a strong POV is challenging but shapes the entire future of the project. Always record the "
   "underlying assumptions so they can be tested. The Business Model Canvas is a complementary Define tool for "
   "reflecting systematically on the business model segment by segment."),
  ("Ideation methods",
   "There are hundreds: brainstorm, mind map, sketch or sketchstorm, storyboard, analogies, provocation, movement, "
   "bodystorm, gamestorming, cheatstorm, crowdstorm, co-creation workshops, prototyping and the creative pause. "
   "Brainstorming leverages the collective thinking of the group by engaging, listening and building on others' ideas. "
   "Rules: set a mission, set up the space, limit the time, don't hang on to any one idea too tightly, defer judgement "
   "and go for volume — 10 ideas are better than 3, and 200 are better than 50. Once the session is complete, collect, "
   "categorise, refine and narrow down using post-it or dot voting, the four categories method, bingo selection, idea "
   "affinity maps or a Now-Wow-How matrix."),
  ("Prototyping and testing in practice",
   "Prototype as if you know you're right, but test as if you know you're wrong. Ask what the cheapest, quickest test "
   "is that would prove or disprove your assumption. Give the test a name and a description, decide where and when you "
   "will run it, and state what metrics will measure success — pre-orders, votes, completion, smiles. Write down which "
   "assumptions this particular prototype tests and the minimum result that would count as success. When planning the "
   "test, let users compare alternatives, show rather than tell, and ask them to talk through their experience."),
  ("Measuring the outcome",
   "Measure design thinking across several categories: traditional KPIs (increased sales, ROI per project and other "
   "financial measures); customer feedback (customer satisfaction, net promoter score, campaign response, usability "
   "metrics); design-thinking activity (number of projects, people trained, coaches trained); and quick results "
   "(concepts finished, projects launched, funded or in development). Move beyond execution-oriented metrics to track "
   "creative behaviours. The three main drivers leading companies to pursue design thinking are: to better understand "
   "customers or end users, to protect business share from disruption and start-ups, and to develop more innovative "
   "methods and team dynamics. Plot concepts on a business-value versus novelty grid — valuable and novel is the target "
   "quadrant."),
 ],
}

for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}")
    p(t["subtitle"])
    p(f"Skills mapped: {t['weighting']}")
    h3("Key concepts")
    bullets([f"{k} — {v}" for k,v in t["concepts"]])
    for hd,body in TOPIC_NOTES[t["num"]]:
        h3(hd); p(body)
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Activity {a['num']} — {a['title']}")
        p(f"Objective: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll produce")
        p(a["build"]+f"   (Ed-tools: {a['services']}.)")
        h3("Step-by-step")
        steps([(instr,cmd) for instr,cmd in a["steps"]])
        h3("You're done when")
        p(a["test"])
        note(f"The full activity brief is in labs/lab-{a['num']:02d}-*.md. "
             f"Keep your output — it is the evidence used in the Practical Performance assessment.")
        rule()

h1("Preparing for the Assessment")
h3("Written Assessment (WA) — Short-Answer Questions")
bullets([
 f"Open book, {C.WA_MINUTES} minutes. You may use these slides, this Learner Guide and approved materials.",
 "Revise the definition of design thinking and why it is human-centred (Topic 1).",
 "Be able to name the traits and mindset of a design thinker, and contrast them with traditional thinking.",
 "Know the three lenses — desirability, viability, feasibility — and what each asks.",
 "Be able to describe at least two real use cases and what changed as a result (Topic 2).",
 "Know all five action phases in order, and the objective of each (Topic 3).",
 "Know the four quadrants of the empathy map and the POV formula (Topic 4).",
 "Be able to list metric categories used to measure design outcomes (Topic 4).",
])
h3("Practical Performance (PP)")
bullets([
 f"Open book, {C.PP_MINUTES} minutes, using the browser-based ed-tools.",
 "You will be asked to perform design-thinking tasks of the same kind as the class activities.",
 "Expect to build or interpret an empathy map, write a POV statement, generate and select ideas, describe a prototype and state how you would measure success.",
 "Your class activity outputs are valid reference material — keep them open.",
 "Write in full sentences and justify your choices; the assessor is looking for reasoning, not just an answer.",
])

h1("Glossary")
gl=[
 ("Design thinking","A human-centred, iterative problem-solving approach that aims to improve people's experience."),
 ("Empathize","The first action phase — understanding the user's experience, situation and emotion by observing, engaging and immersing."),
 ("Define","The second action phase — synthesising findings into an actionable problem statement."),
 ("Ideate","The third action phase — generating a wide variety and large quantity of possible solutions."),
 ("Prototype","The fourth action phase — building a cheap, fast, tangible representation of an idea."),
 ("Test","The fifth action phase — putting the prototype in users' hands to gather feedback and refine."),
 ("Empathy map","A four-quadrant canvas (Says, Thinks, Does, Feels) capturing a user as a whole."),
 ("Persona","A named, fictional but evidence-based representation of a user segment, with goals and frustrations."),
 ("Journey map","A chart of every touchpoint a user has with a service over time, exposing pain points."),
 ("Point of View (POV)","A problem statement in the form: [USER] needs a way to [NEED] because [INSIGHT]."),
 ("How Might We (HMW)","A reframing of a POV as an open question broad enough to allow many answers."),
 ("Divergent thinking","Deliberately generating many options before evaluating any of them."),
 ("Convergent thinking","Narrowing a wide set of options down to the strongest few."),
 ("Low-fidelity prototype","A basic, quick model using simple materials — a sketch, storyboard or paper mock-up."),
 ("High-fidelity prototype","A prototype close to the finished product, used later to test detail."),
 ("Desirability / Viability / Feasibility","The three lenses of organisational design thinking: user desire, business viability, technical feasibility."),
 ("Beginner's mindset","Consciously setting aside your own assumptions and expertise when observing users."),
 ("Net Promoter Score (NPS)","A customer-feedback metric measuring how likely users are to recommend a product or service."),
 ("RACI","A responsibility matrix assigning who is Responsible, Accountable, Consulted and Informed."),
 ("TRAQOM","The mandatory SSG training-quality survey learners complete at the end of a WSQ course."),
]
B.append(("dl",gl))

h1("References and Further Reading")
bullets([
 "Stanford d.school — An Introduction to Design Thinking Process Guide, dschool.stanford.edu",
 "Tim Brown, IDEO — Change by Design and ideo.com/case-study",
 "IBM Design Thinking — ibm.com/design/thinking/",
 "Design Management Institute — dmi.org/page/2015DVIandOTW (design value index)",
 "This is Design Thinking — thisisdesignthinking.net/category/cases/",
 "IDEO Design Kit case studies — designkit.org/case-studies",
 "Harvard Business Review — 'Better Service, Faster: A Design Thinking Case Study', hbr.org/2016/01",
 "Nielsen Norman Group — empathy mapping and journey mapping guidance, nngroup.com",
])

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **TSC:** {C.TSC_TITLE} ({C.TSC_CODE})  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("14.0","1 January 2024","Previous release of the Learner Guide, issued under the superseded course reference.",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  f"Re-issued under WSQ course code {C.COURSE_CODE}. Content carried over from the v14 master trainer deck and restructured to the current Tertiary Infotech WSQ house standard. Added {len(ACT)} step-by-step hands-on activities built on the browser-based ed-tools, expanded topic notes, an assessment-preparation section, a glossary and references.",C.TRAINER),
])
prodoc.add_toc(doc)

def _decimal_abstract_id():
    """Return the abstractNumId of a DECIMAL numbering definition, creating one if the
    document has none. Picking the first <w:num> blindly can land on a BULLET definition,
    which silently renders every 'step' as a bullet — so match on numFmt explicitly."""
    numbering=doc.part.numbering_part.element
    for abs_el in numbering.findall(qn("w:abstractNum")):
        lvl=abs_el.find(qn("w:lvl"))
        if lvl is None: continue
        fmt=lvl.find(qn("w:numFmt"))
        if fmt is not None and fmt.get(qn("w:val"))=="decimal":
            return abs_el.get(qn("w:abstractNumId"))
    # none present — build a minimal decimal definition
    used=[int(a.get(qn("w:abstractNumId"))) for a in numbering.findall(qn("w:abstractNum"))] or [-1]
    new_abs=str(max(used)+1)
    abs_el=OxmlElement("w:abstractNum"); abs_el.set(qn("w:abstractNumId"),new_abs)
    lvl=OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"),"0")
    for tag,val in (("w:start","1"),("w:numFmt","decimal"),("w:lvlText","%1."),("w:lvlJc","left")):
        el=OxmlElement(tag); el.set(qn("w:val"),val); lvl.append(el)
    pPr=OxmlElement("w:pPr"); ind=OxmlElement("w:ind")
    ind.set(qn("w:left"),"720"); ind.set(qn("w:hanging"),"360"); pPr.append(ind); lvl.append(pPr)
    abs_el.append(lvl); numbering.insert(0,abs_el)
    return new_abs

_DECIMAL_ABS=None
def _new_numbering_id():
    """Create a FRESH concrete <w:num> bound to the DECIMAL abstract definition, so each
    activity's step list is numbered and restarts at 1."""
    global _DECIMAL_ABS
    numbering=doc.part.numbering_part.element
    if _DECIMAL_ABS is None:
        _DECIMAL_ABS=_decimal_abstract_id()
    used=[int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))] or [0]
    new_id=max(used)+1
    num=OxmlElement("w:num"); num.set(qn("w:numId"),str(new_id))
    aid=OxmlElement("w:abstractNumId"); aid.set(qn("w:val"),str(_DECIMAL_ABS)); num.append(aid)
    # Sharing one abstractNum means the counter is shared too — Word and LibreOffice will
    # keep counting 1,2,3… straight across every activity. An explicit lvlOverride with
    # startOverride=1 forces THIS list to begin again at 1.
    ov=OxmlElement("w:lvlOverride"); ov.set(qn("w:ilvl"),"0")
    st=OxmlElement("w:startOverride"); st.set(qn("w:val"),"1"); ov.append(st)
    num.append(ov)
    numbering.append(num)
    return new_id

def _set_num(para,num_id):
    pPr=para._p.get_or_add_pPr()
    numPr=pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr=OxmlElement("w:numPr"); pPr.append(numPr)
    for tag in ("w:ilvl","w:numId"):
        el=numPr.find(qn(tag))
        if el is not None: numPr.remove(el)
    ilvl=OxmlElement("w:ilvl"); ilvl.set(qn("w:val"),"0"); numPr.append(ilvl)
    nid=OxmlElement("w:numId"); nid.set(qn("w:val"),str(num_id)); numPr.append(nid)

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        num_id=_new_numbering_id()          # fresh list → restarts at 1
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph(style="List Number"); para.add_run(instr)
            _set_num(para,num_id)
            if cmd:
                cp=doc.add_paragraph(); r=cp.add_run(cmd)
                r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
